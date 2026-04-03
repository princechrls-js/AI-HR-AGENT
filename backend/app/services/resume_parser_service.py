import fitz  # PyMuPDF
from docx import Document
from app.services.ocr_service import ocr_service
from app.utils.text_cleaner import text_cleaner
from app.core.logging import logger
import os
import re

class ResumeParserService:
    def extract_text(self, file_path: str) -> str:
        if not file_path or not os.path.exists(file_path):
            logger.error(f"Resume file not found at: {file_path}")
            return ""
            
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self._extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file extension: {extension}")
            return ""

    def _extract_from_pdf(self, pdf_path: str) -> str:
        """
        Enhanced PDF extraction using structural block analysis.
        Uses PyMuPDF's block-level extraction to maintain document structure,
        properly handle multi-column layouts, and preserve section boundaries.
        """
        try:
            doc = fitz.open(pdf_path)
            all_text_parts = []

            for page_num, page in enumerate(doc):
                # Method 1: Try block-based extraction (handles columns correctly)
                blocks = page.get_text("blocks")
                
                if blocks:
                    # Sort blocks by vertical position (y0), then horizontal (x0)
                    # This ensures proper reading order even in multi-column layouts
                    sorted_blocks = sorted(blocks, key=lambda b: (round(b[1] / 15) * 15, b[0]))
                    
                    page_text_parts = []
                    prev_y = -1
                    
                    for block in sorted_blocks:
                        # block format: (x0, y0, x1, y1, text, block_no, block_type)
                        if block[6] != 0:  # Skip image blocks (type 1)
                            continue
                        
                        text = block[4].strip()
                        if not text:
                            continue
                        
                        current_y = round(block[1])
                        
                        # Add section break if there's a significant vertical gap
                        if prev_y >= 0 and (current_y - prev_y) > 25:
                            page_text_parts.append("")  # Empty line as section separator
                        
                        page_text_parts.append(text)
                        prev_y = round(block[3])  # bottom of current block
                    
                    page_text = "\n".join(page_text_parts)
                else:
                    # Fallback to simple text extraction
                    page_text = page.get_text()
                
                if page_text.strip():
                    all_text_parts.append(page_text)

            doc.close()
            
            combined_text = "\n\n".join(all_text_parts)

            # If extraction was poor (scanned PDF), fall back to OCR
            if len(combined_text.strip()) < 50:
                logger.info(f"PDF text extraction poor for {pdf_path}. Triggering OCR fallback.")
                try:
                    combined_text = ocr_service.extract_text_from_scanned_pdf(pdf_path)
                except Exception as ocr_err:
                    logger.warning(f"OCR fallback also failed: {ocr_err}")
            
            # Post-processing: clean up common PDF artifacts
            cleaned = self._post_process_resume_text(combined_text)
            return text_cleaner.clean_text(cleaned)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return ""

    def _post_process_resume_text(self, text: str) -> str:
        """
        Clean up common PDF extraction artifacts for better AI analysis.
        """
        # Remove excessive whitespace within lines
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page headers/footers (common patterns)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\bpage\s*\d+\b', '', text, flags=re.IGNORECASE)
        
        # Normalize bullet points
        text = re.sub(r'[•●◦▪▸►]', '•', text)
        text = re.sub(r'^\s*[-–—]\s', '• ', text, flags=re.MULTILINE)
        
        # Remove excessive blank lines
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        
        # Fix common ligature issues
        text = text.replace('ﬁ', 'fi').replace('ﬂ', 'fl').replace('ﬀ', 'ff')
        
        return text.strip()

    def _extract_from_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)
            
            # Also extract from tables (many resumes use tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        parts.append(" | ".join(row_text))
            
            text = "\n".join(parts)
            return text_cleaner.clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return ""

resume_parser_service = ResumeParserService()
